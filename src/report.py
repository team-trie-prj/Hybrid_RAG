"""보고서 자동 생성 모듈: 검색 결과 + 통계를 요약해 .docx로 산출 (도메인 중립).

질의 의도(도메인·지역)에 맞춰 검색 결과 표 + 관련 통계 표를 자동 구성한다.
"""
from __future__ import annotations

import os
from datetime import datetime, timezone

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt

import agent
from retrieve import Index

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
TYPE_KO = {"pothole": "포트홀", "accident_stat": "사고통계", "road_infra": "도로인프라",
           "vision_detection": "비전탐지", "air_quality": "대기질", "fire_stat": "화재",
           "population": "인구", "facility": "시설"}


def _font(doc, name="맑은 고딕"):
    s = doc.styles["Normal"]; s.font.name = name; s.font.size = Pt(10.5)
    s._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def _k(par, name="맑은 고딕"):
    for run in par.runs:
        run.font.name = name
        run._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def _plan(question: str):
    """질의 의도 → (filters, metric, label, group_by)."""
    intent = agent.analyze_intent(question)
    flt = {}
    if intent["domain"]:
        flt["domain"] = intent["domain"]
    if intent["sigungu_cd"]:
        flt["sigungu_cd"] = intent["sigungu_cd"]
    elif intent["sido_cd"]:
        flt["sido_cd"] = intent["sido_cd"]
    metric, label = agent.DOMAIN_METRIC.get(intent["domain"] or "교통안전", ("count", "건수"))
    group_by = "surface" if "노면" in question else (
        "roadtype" if "도로종류" in question else "region_name")
    return intent, flt, metric, label, group_by


def _collect(question: str, idx: Index):
    """질의 → (검색결과, 통계, 메타). docx/markdown 보고서가 공유하는 데이터 수집부."""
    intent, flt, metric, label, group_by = _plan(question)
    hits = idx.hybrid_search(question, filters=flt, k=6)
    agg_flt = dict(flt)
    if group_by == "region_name":
        agg_flt.pop("sido_cd", None); agg_flt.pop("sigungu_cd", None)
    stats = idx.stats_query(metric, group_by, agg_flt)[:8]
    region = intent["region"] or "전국"
    domain = intent["domain"] or "공공데이터"
    return intent, hits, stats, metric, label, group_by, region, domain


def _summary_text(question, hits, stats, metric, label, region, domain) -> str:
    """FR-RPT-001 자동 요약: 실제 검색·집계 결과만 반영(원문 외 사실 미생성)."""
    top = stats[0] if stats else None
    s = (f"'{question}' 질의에 대해 {region} {domain} 도메인에서 관련 문서 {len(hits)}건을 검색했다. "
         f"{label} 기준 집계 결과 ")
    s += (f"'{top['group']}'이(가) {int(top[metric]):,}로 가장 높다. " if top else "집계 대상이 없다. ")
    s += "사내 비전 데이터와 공공데이터를 연계해 의사결정 근거로 활용할 수 있다."
    return s


def generate_report(question: str, idx: Index | None = None,
                    out_path: str | None = None, summary: str | None = None) -> str:
    idx = idx or Index()
    out_path = out_path or os.path.join(ROOT, "report_output.docx")
    intent, hits, stats, metric, label, group_by, region, domain = _collect(question, idx)
    if not summary:
        summary = _summary_text(question, hits, stats, metric, label, region, domain)

    doc = docx.Document(); _font(doc)
    t = doc.add_heading(f"{region} {domain} 분석 보고서", 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER; _k(t)
    meta = doc.add_paragraph(); meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"질의: {question}\n작성일: "
                 f"{datetime.now(timezone.utc).astimezone().strftime('%Y-%m-%d %H:%M')} · "
                 f"하이브리드 RAG 자동 생성"); _k(meta)

    h = doc.add_heading("1. 요약", level=1); _k(h)
    _k(doc.add_paragraph(summary))

    h = doc.add_heading("2. 검색 결과 (하이브리드 검색 + 재정렬)", level=1); _k(h)
    tb = doc.add_table(rows=1, cols=4); tb.style = "Light Grid Accent 1"
    for c, x in zip(tb.rows[0].cells, ("종류", "제목", "지역/도로", "출처ID")):
        c.paragraphs[0].add_run(x).bold = True; _k(c.paragraphs[0])
    for hh in hits:
        cs = tb.add_row().cells
        loc = hh.get("region_name") or hh.get("road_name") or "-"
        for cell, v in zip(cs, (TYPE_KO.get(hh["doc_type"], hh["doc_type"]), hh["title"],
                                str(loc), hh["provenance"].get("dataset_id", "-"))):
            cell.paragraphs[0].add_run(str(v)); _k(cell.paragraphs[0])

    if stats:
        h = doc.add_heading(f"3. 통계 집계 — {label} ({group_by})", level=1); _k(h)
        st = doc.add_table(rows=1, cols=2); st.style = "Light Grid Accent 1"
        for c, x in zip(st.rows[0].cells, ("구분", label)):
            c.paragraphs[0].add_run(x).bold = True; _k(c.paragraphs[0])
        for s in stats:
            cs = st.add_row().cells
            cs[0].paragraphs[0].add_run(str(s["group"])); _k(cs[0].paragraphs[0])
            cs[1].paragraphs[0].add_run(f"{int(s[metric]):,}"); _k(cs[1].paragraphs[0])

    h = doc.add_heading("4. 데이터 출처", level=1); _k(h)
    seen = set()
    for hh in hits:
        url = hh["provenance"].get("url", "")
        if url and url not in seen:
            seen.add(url)
            _k(doc.add_paragraph(f"{hh['provenance'].get('dataset_id','-')} — {url}",
                                 style="List Bullet"))
    doc.save(out_path)
    return out_path


def build_markdown(question: str, idx: Index | None = None,
                   summary: str | None = None) -> str:
    """FR-RPT-002/003: 제목·요약·본문·출처 구조의 마크다운 보고서 텍스트를 생성한다.

    - FR-RPT-001 요약: 실제 검색·집계 결과 기반(원문 외 사실 미생성).
    - FR-RPT-003 근거: 검색 결과 표와 출처(dataset_id·URL)를 포함.
    """
    idx = idx or Index()
    intent, hits, stats, metric, label, group_by, region, domain = _collect(question, idx)
    summary = summary or _summary_text(question, hits, stats, metric, label, region, domain)
    ts = datetime.now(timezone.utc).astimezone().strftime("%Y-%m-%d %H:%M")

    lines: list[str] = []
    lines.append(f"# {region} {domain} 분석 보고서")
    lines.append("")
    lines.append(f"> 질의: **{question}**  ")
    lines.append(f"> 작성일: {ts} · 하이브리드 RAG 자동 생성  ")
    lines.append(f"> 처리 모듈: {intent['intent']} → {intent['module']}")
    lines.append("")
    lines.append("## 1. 요약")
    lines.append("")
    lines.append(summary)
    lines.append("")
    lines.append("## 2. 검색 결과 (하이브리드 검색 + 재정렬)")
    lines.append("")
    if hits:
        lines.append("| 종류 | 제목 | 지역/도로 | 유사도 | 출처ID |")
        lines.append("| --- | --- | --- | --- | --- |")
        for h in hits:
            loc = h.get("region_name") or h.get("road_name") or "-"
            lines.append(f"| {TYPE_KO.get(h['doc_type'], h['doc_type'])} | {h['title']} | "
                         f"{loc} | {h.get('_similarity', '-')} | "
                         f"{h['provenance'].get('dataset_id', '-')} |")
    else:
        lines.append("_관련 문서를 찾을 수 없습니다._")
    lines.append("")
    if stats:
        lines.append(f"## 3. 통계 집계 — {label} ({group_by})")
        lines.append("")
        lines.append(f"| 구분 | {label} |")
        lines.append("| --- | ---: |")
        for s in stats:
            lines.append(f"| {s['group']} | {int(s[metric]):,} |")
        lines.append("")
    lines.append("## 4. 데이터 출처")
    lines.append("")
    seen = set()
    for h in hits:
        url = h["provenance"].get("url", "")
        if url and url not in seen:
            seen.add(url)
            lines.append(f"- {h['provenance'].get('dataset_id', '-')} — {url}")
    if not seen:
        lines.append("- (출처 없음)")
    lines.append("")
    return "\n".join(lines)


def generate_markdown_report(question: str, idx: Index | None = None,
                             out_path: str | None = None,
                             summary: str | None = None) -> tuple[str, str]:
    """FR-RPT-004 내보내기: 마크다운 보고서를 .md 파일로 저장. (텍스트, 경로) 반환."""
    out_path = out_path or os.path.join(ROOT, "report_output.md")
    md = build_markdown(question, idx, summary)
    with open(out_path, "w", encoding="utf-8") as f:
        f.write(md)
    return md, out_path


if __name__ == "__main__":
    import sys
    args = sys.argv[1:]
    as_md = "--md" in args
    args = [a for a in args if a != "--md"]
    q = " ".join(args) or "대전 노면상태별 교통사고 통계를 요약해줘"
    if as_md:
        md, path = generate_markdown_report(q)
        print(md)
        print(f"\n[저장됨] {path}")
    else:
        print(f"보고서 생성 완료 → {generate_report(q)}")
